from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
 
class ResumeTemplate:
    
    def CreateResume(self,data):
        # Create a new word document
        document = Document()
        section = document.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_height = Inches(11.0)
        section.page_width = Inches(8.5)
        section.header_distance = Inches(0.5)

        # Add a heading to the document with custom font and size
        header = document.add_heading(data['Name'], 0)
        header.style.font.name = 'Arial'
        header.style.font.size = Pt(20)
        header.style.font.color.rgb = RGBColor(1, 0, 0)

        # Add your name, address, contact information with custom font and size
        document.add_paragraph(data['Address'], style='Normal').style.font.name = 'Arial'
        document.add_paragraph(data['Email'], style='Normal').style.font.name = 'Arial'
        document.add_paragraph(data['Phone'], style='Normal').style.font.name = 'Arial'
        document.add_paragraph(data['LinkedIn'], style='Normal').style.font.name = 'Arial'

        # Add a section for your objective
        document.add_heading('Objective', level=1)
        document.add_paragraph(data['Objective'], style='Normal').style.font.name = 'Arial'

        # Add a section for your work experience
        document.add_heading('Work Experience', level=1)
        # Add your work experience, job title, and duration as bullet points
        experiences = data['experiences']
        for exp in experiences:
            company = exp['company']
            title = exp['title']
            duration = exp['duration']
            description = exp['description']
            
            # Create a heading for the job title and company
            heading = document.add_heading(company + ', ' + title, level=2)
            heading.style.font.name = 'Arial'
            heading.style.font.size = Pt(14)
            heading.style.font.bold = True
            
            # Add the duration and description as a paragraph
            paragraph = document.add_paragraph(duration + '\n' + description, style='Normal')
            paragraph.style.font.name = 'Arial'


    # Add a section for your education
        document.add_heading('Education', level=1)
        education = data["education"]
        for edu in education:
            school = edu["school"]
            degree = edu["degree"]
            start_date = edu["start_date"]
            end_date = edu["end_date"]
            city = edu["city"]
            description = edu.get("description", "")

            p = document.add_paragraph( )
            p.add_run(school).bold = True
            p.add_run(', ' + city)
            p.add_run(' | ' + degree)

            p.add_run('\n' + start_date + ' - ' + end_date)
            
            if description:
                p.add_run('\n' + description).italic = True


        # Add a section for your technical skills
        document.add_heading('Skills', level=1)
        skills = data.get("Skills", "")
        if skills:
            document.add_paragraph(skills, style='Normal').style.font.name = 'Arial'

        # Add a section for your achievements
        document.add_heading('Achievements', level=1)
        achievements = data.get("Achievement", "")
        if achievements:
            document.add_paragraph(achievements, style='Normal').style.font.name = 'Arial'

        return document

