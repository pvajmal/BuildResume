import json
from docx import Document
from docx2pdf import convert
import os
import pdfkit
import docx

class CreateResume:
    def create_resume(self, template_file, data):
        # Open the template file
        doc = Document(template_file)

        # Define mapping of template placeholder and data key
        mapping = {
            "{{Name}}": "Name",
            "{{Objective}}": "Objective",
            "{{Address}}": "Address",
            "{{Phone}}": "Phone",
            "{{Email}}": "Email",
            "{{JOBTITLE}}": "JOBTITLE",
            "{{COMPANY}}": "COMPANY",
            "{{ExpPlace}}": "ExpPlace",
            "{{ExpDuration}}": "ExpDuration",
            "{{job description}}": "job description",
            "{{Technical Skills}}": "Technical Skills",
            "{{SSkills}}": "SSkills",
            "{{LinkedIn}}" : "LinkedIn",
            "{{College}}": "College",
            "{{Degree}}": "Degree",
            "{{CollegePlace}}": "CollegePlace",
            "{{CGPA}}": "CGPA",
            "{{CollegeDetails}}": "CollegeDetails",
            "{{CollegeDuration}}": "CollegeDuration",
            "{{Achieve}}": "Achievement"
        }


        # Iterate through the document's paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                for key, value in mapping.items():
                    if key in run.text:

                        run.text = run.text.replace(key, data.get(value, ''))

        # Save the document with a new file name
        file_path = os.path.join('Output', f"{data['Name']}_Resume.docx")
        doc.save(file_path)


    # Convert to pdf
    def convert_docx_to_pdf(self, docx_file, pdf_file):
        doc = docx.Document(docx_file)
        pdfkit.from_string(doc.text, pdf_file)

    def save_data(self, data):
        with open("data.json", "w") as f:
            json.dump(data, f)


    def load_data(self):
        try:
            with open("data.json", "r") as f:
                return json.load(f)
        except FileNotFoundError:
            return {}